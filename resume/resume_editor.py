from __future__ import annotations

import os
import re
import shutil
import subprocess
from datetime import datetime
from typing import List, Optional

from config import settings


async def edit_resume(
    original_bullets: List[str],
    tailored_bullets: List[str],
    company: str,
    job_title: str,
) -> str:
    """
    Edit the base resume to swap in tailored bullets.
    Returns path to the new tailored resume file.
    """
    os.makedirs(settings.tailored_resume_dir, exist_ok=True)

    date_str = datetime.now().strftime("%Y%m%d")
    safe_company = re.sub(r"[^\w]", "_", company)[:30]

    if settings.resume_format == "latex":
        return _edit_latex(original_bullets, tailored_bullets, safe_company, date_str)
    else:
        return _edit_docx(original_bullets, tailored_bullets, safe_company, date_str)


def _edit_docx(
    original_bullets: List[str],
    tailored_bullets: List[str],
    company: str,
    date_str: str,
) -> str:
    """Use python-docx to find and replace bullet text."""
    from docx import Document

    src = settings.base_resume_docx
    if not os.path.exists(src):
        raise FileNotFoundError(
            f"Base resume not found: {src}. "
            "Please place your resume at resume/base_resume.docx"
        )

    out_path = os.path.join(settings.tailored_resume_dir, f"resume_{company}_{date_str}.docx")
    shutil.copy2(src, out_path)

    doc = Document(out_path)
    bullet_map = dict(zip(original_bullets, tailored_bullets))

    for paragraph in doc.paragraphs:
        para_text = paragraph.text.strip()
        if para_text in bullet_map:
            new_text = bullet_map[para_text]
            # Preserve runs/formatting — replace text in first run
            if paragraph.runs:
                # Clear all runs except the first
                for run in paragraph.runs[1:]:
                    run.text = ""
                paragraph.runs[0].text = new_text
            else:
                paragraph.text = new_text

    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para_text = para.text.strip()
                    if para_text in bullet_map:
                        if para.runs:
                            for run in para.runs[1:]:
                                run.text = ""
                            para.runs[0].text = bullet_map[para_text]

    doc.save(out_path)
    print(f"[ResumeEditor] Saved tailored docx: {out_path}")
    return out_path


def _edit_latex(
    original_bullets: List[str],
    tailored_bullets: List[str],
    company: str,
    date_str: str,
) -> str:
    """Edit LaTeX source and compile to PDF."""
    src = settings.base_resume_tex
    if not os.path.exists(src):
        raise FileNotFoundError(
            f"Base resume not found: {src}. "
            "Please place your resume at resume/base_resume.tex"
        )

    tex_out = os.path.join(settings.tailored_resume_dir, f"resume_{company}_{date_str}.tex")
    shutil.copy2(src, tex_out)

    with open(tex_out, "r", encoding="utf-8") as f:
        content = f.read()

    for orig, new in zip(original_bullets, tailored_bullets):
        # Escape special LaTeX chars in the search string for safe replacement
        escaped_orig = re.escape(orig)
        content = re.sub(escaped_orig, _escape_latex(new), content)

    with open(tex_out, "w", encoding="utf-8") as f:
        f.write(content)

    # Compile to PDF
    pdf_path = _compile_latex(tex_out, settings.tailored_resume_dir)
    return pdf_path if pdf_path else tex_out


def _compile_latex(tex_path: str, output_dir: str) -> Optional[str]:
    """Run pdflatex to compile .tex → .pdf. Returns PDF path or None."""
    if not shutil.which("pdflatex"):
        print("[ResumeEditor] pdflatex not found. Returning .tex path only.")
        return None

    try:
        result = subprocess.run(
            [
                "pdflatex",
                "-interaction=nonstopmode",
                f"-output-directory={output_dir}",
                tex_path,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode == 0:
            pdf_path = tex_path.replace(".tex", ".pdf")
            if os.path.exists(pdf_path):
                print(f"[ResumeEditor] Compiled PDF: {pdf_path}")
                return pdf_path
    except subprocess.TimeoutExpired:
        print("[ResumeEditor] pdflatex timed out.")
    except Exception as exc:
        print(f"[ResumeEditor] LaTeX compile error: {exc}")

    return None


def _escape_latex(text: str) -> str:
    """Escape special LaTeX characters in a replacement string."""
    replacements = {
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
        "\\": r"\textbackslash{}",
    }
    for char, escaped in replacements.items():
        text = text.replace(char, escaped)
    return text


def extract_bullets_from_docx(path: str) -> List[str]:
    """
    Extract all bullet/list paragraph text from a .docx file.
    Returns list of non-empty bullet strings.
    """
    from docx import Document
    from docx.oxml.ns import qn

    if not os.path.exists(path):
        return []

    doc = Document(path)
    bullets = []

    for para in doc.paragraphs:
        # Check for list-style paragraphs
        style_name = para.style.name.lower() if para.style else ""
        is_list = (
            "list" in style_name
            or "bullet" in style_name
            or para.text.strip().startswith(("•", "-", "●", "*", "◦"))
        )

        # Also check numPr (numbered/bulleted via Word's list formatting)
        try:
            num_pr = para._element.find(qn("w:numPr"))
            if num_pr is not None:
                is_list = True
        except Exception:
            pass

        text = para.text.strip()
        if is_list and text and len(text) > 10:
            # Clean leading bullet chars
            text = text.lstrip("•-●*◦ \t")
            bullets.append(text)

    return bullets


def extract_full_text_from_docx(path: str) -> str:
    """Extract all text from a .docx file as a single string."""
    from docx import Document

    if not os.path.exists(path):
        return ""
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ── PDF support ───────────────────────────────────────────────────────────────

def extract_full_text_from_pdf(path: str) -> str:
    """
    Extract all text from a PDF resume.
    Tries pdfminer.six first (best layout preservation),
    falls back to pypdf if pdfminer is not installed.
    """
    if not os.path.exists(path):
        return ""

    # Strategy 1: pdfminer.six
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(path)
        if text and text.strip():
            return text.strip()
    except ImportError:
        pass
    except Exception as exc:
        print(f"[ResumeEditor] pdfminer error: {exc}")

    # Strategy 2: pypdf
    try:
        import pypdf
        reader = pypdf.PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages).strip()
        if text:
            return text
    except ImportError:
        pass
    except Exception as exc:
        print(f"[ResumeEditor] pypdf error: {exc}")

    print("[ResumeEditor] WARNING: Could not extract PDF text. Install pdfminer.six: pip install pdfminer.six")
    return ""


def extract_bullets_from_pdf(path: str) -> List[str]:
    """
    Extract bullet-like lines from a PDF resume.

    Strategy (in order):
      1. Lines that start with a recognised bullet glyph (Unicode-aware).
      2. Lines that start with a Markdown-style bullet (-, *, +, numbered).
      3. Lines that match a broad action-verb pattern.
      4. Lines longer than 40 chars that look like achievement sentences
         (start with capital, contain metrics or tech keywords).

    PDF text extraction collapses formatting, so we also handle:
      - Multi-line bullets that are joined by pdfminer on one line.
      - Spurious short header/footer lines (skipped if < 15 chars).
      - Unicode ligature / hyphen variants (\u2013, \u2014, etc.).
    """
    full_text = extract_full_text_from_pdf(path)
    if not full_text:
        return []

    # Bullet glyphs pdfminer preserves from PDF encoding
    BULLET_GLYPHS = frozenset(
        "\u2022\u2023\u2043\u204c\u204d\u2219\u25aa\u25ab"
        "\u25b8\u25cf\u25e6\u2605\u2606\u29bf\u30fb"
        "\u2212\u2013\u2014"  # en-dash / em-dash also used as bullets
        "\u25ba\u2192\u27a4\u2714\u2718"
    )
    ASCII_BULLETS = frozenset("-", "*", "+", "\u00b7", "\u00bb", "\u00b9")
    ALL_BULLETS = BULLET_GLYPHS | ASCII_BULLETS

    # Broad action-verb list (past tense + gerunds)
    VERB_PATTERN = re.compile(
        r"^(Developed|Built|Designed|Implemented|Led|Created|Improved|Reduced|"
        r"Achieved|Deployed|Trained|Optimized|Automated|Researched|Analysed|"
        r"Analyzed|Managed|Delivered|Increased|Decreased|Collaborated|Published|"
        r"Fine.?tuned|Engineered|Contributed|Integrated|Migrated|Established|"
        r"Spearheaded|Architected|Streamlined|Accelerated|Launched|Maintained|"
        r"Monitored|Configured|Supported|Designed|Refactored|Documented|"
        r"Evaluated|Investigated|Prototyped|Demonstrated|Presented|Proposed|"
        r"Executed|Facilitated|Coordinated|Produced|Resolved|Identified|"
        r"Developed|Gathered|Generated|Reviewed|Tested|Validated|Implemented)",
        re.IGNORECASE,
    )

    # Numbered list pattern:  "1. ... "  "(a) ..."  "a) ..."
    NUMBERED_RE = re.compile(r"^(\d{1,2}[.)\]]|[a-z][.)\]])\s+", re.IGNORECASE)

    bullets: List[str] = []
    seen: set = set()

    for line in full_text.splitlines():
        line = line.rstrip()
        if not line:
            continue

        stripped = line.lstrip()
        if len(stripped) < 15:          # skip section headers, page numbers, etc.
            continue

        # Detect and strip leading bullet glyph
        first_char = stripped[0] if stripped else ""
        is_bullet = False
        cleaned = stripped

        if first_char in ALL_BULLETS:
            is_bullet = True
            cleaned = stripped[1:].lstrip()
        elif NUMBERED_RE.match(stripped):
            is_bullet = True
            cleaned = NUMBERED_RE.sub("", stripped).lstrip()
        elif stripped.startswith(("- ", "* ", "+ ")):
            # Explicit space after ASCII bullet prevents false positives
            is_bullet = True
            cleaned = stripped[2:].lstrip()

        if not cleaned or len(cleaned) < 15:
            continue

        norm = cleaned.strip()

        if is_bullet:
            if norm not in seen:
                seen.add(norm)
                bullets.append(norm)
        elif VERB_PATTERN.match(norm):
            if norm not in seen:
                seen.add(norm)
                bullets.append(norm)
        elif (
            len(norm) >= 40
            and norm[0].isupper()
            and any(kw in norm.lower() for kw in (
                "%", "x ", "k ", "m ", "million", "billion",
                "python", "model", "system", "api", "data", "ml",
                "accuracy", "latency", "throughput", "reduced", "improved",
            ))
        ):
            if norm not in seen:
                seen.add(norm)
                bullets.append(norm)

    return bullets


def detect_resume_format(path: str) -> str:
    """Return 'docx', 'pdf', or 'tex' based on file extension."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return "pdf"
    if ext in (".tex",):
        return "latex"
    return "docx"


def load_resume(path: str):
    """
    Universal resume loader.
    Returns (full_text: str, bullets: List[str]) for any supported format.
    """
    fmt = detect_resume_format(path)
    if fmt == "pdf":
        return (
            extract_full_text_from_pdf(path),
            extract_bullets_from_pdf(path),
        )
    if fmt == "latex":
        text = open(path, encoding="utf-8").read() if os.path.exists(path) else ""
        # Strip LaTeX commands for plain text
        text = re.sub(r"\\[a-zA-Z]+\{([^}]*)\}", r"\1", text)
        text = re.sub(r"\\[a-zA-Z]+", " ", text)
        bullets = [
            l.strip() for l in text.splitlines()
            if l.strip().startswith(r"\item") or len(l.strip()) > 20
        ]
        return text, bullets
    # Default: docx
    return (
        extract_full_text_from_docx(path),
        extract_bullets_from_docx(path),
    )
